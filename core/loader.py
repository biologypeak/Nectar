"""
core/loader.py — Universal Document Loader for Nectar
======================================================
Converts any supported format to a list of LangChain Document objects
(same interface as PyPDFLoader). Every format is normalized to Markdown
before reaching the chunking/embedding pipeline.

Conversion strategy per family:
  PDF          → PyPDFLoader  (page-level split, preserves metadata)
  DOCX/ODT/RTF → extracted structure → Markdown headings/tables/bold
  DOC          → LibreOffice headless → DOCX → same as above
  HTML/XML     → BeautifulSoup → markdownify (fallback: .get_text())
  RST          → docutils → HTML → markdownify (fallback: regex strip)
  LaTeX        → pylatexenc LatexNodes2Text
  EPUB         → ebooklib chapters → BeautifulSoup → markdownify
  CSV/TSV      → pandas → Markdown table (200 rows per page)
  JSON/JSONL   → pretty-printed code block
  YAML         → pretty-printed code block
  Code/text    → raw text (code wrapped in fenced block)
"""

from __future__ import annotations

import json
import os
import re
import tempfile
import unicodedata
from pathlib import Path

from langchain_core.documents import Document

# ── Supported extensions ─────────────────────────────────────────────────────
SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({
    # Plain text / logs
    ".txt", ".log",
    # Office
    ".pdf", ".docx", ".doc", ".odt", ".rtf",
    # Markdown / markup
    ".md", ".rst", ".tex", ".html", ".htm", ".xml",
    # Structured data
    ".json", ".jsonl", ".yaml", ".yml", ".csv", ".tsv",
    # E-book
    ".epub",
    # Source code
    ".py", ".js", ".c", ".cpp", ".sh",
})

# Extensions that should NOT be text-cleaned (normalizing would corrupt them)
_STRUCTURED_EXTENSIONS: frozenset[str] = frozenset({
    ".json", ".jsonl", ".yaml", ".yml", ".csv", ".tsv",
    ".py", ".js", ".c", ".cpp", ".sh",
})

EXTENSION_LABELS: dict[str, str] = {
    ".txt": "Plain Text",      ".log": "Log File",
    ".pdf": "PDF",             ".docx": "Word Document",
    ".doc": "Word (Legacy)",   ".odt": "OpenDocument Text",
    ".rtf": "Rich Text",       ".md": "Markdown",
    ".rst": "reStructuredText",".tex": "LaTeX",
    ".html": "HTML",           ".htm": "HTML",
    ".xml": "XML",             ".json": "JSON",
    ".jsonl": "JSON Lines",    ".yaml": "YAML",
    ".yml": "YAML",            ".csv": "CSV",
    ".tsv": "TSV",             ".epub": "EPUB",
    ".py": "Python",           ".js": "JavaScript",
    ".c": "C Source",          ".cpp": "C++ Source",
    ".sh": "Shell Script",
}


def is_structured(filepath: str) -> bool:
    """Return True for code/data formats where text-cleaning must be skipped."""
    return Path(filepath).suffix.lower() in _STRUCTURED_EXTENSIONS


# ── Public entry point ────────────────────────────────────────────────────────
def load_document(filepath: str) -> list[Document]:
    """
    Load *filepath* and return a list of LangChain Documents with Markdown
    page_content. Each Document carries metadata: source, page.
    Raises ValueError for unsupported extensions.
    """
    ext = Path(filepath).suffix.lower()
    dispatch: dict[str, callable] = {
        ".pdf":   _load_pdf,
        ".docx":  _load_docx,
        ".doc":   _load_doc,
        ".odt":   _load_odt,
        ".rtf":   _load_rtf,
        ".epub":  _load_epub,
        ".html":  _load_html,
        ".htm":   _load_html,
        ".xml":   _load_xml,
        ".csv":   _load_csv,
        ".tsv":   _load_tsv,
        ".json":  _load_json,
        ".jsonl": _load_jsonl,
        ".yaml":  _load_yaml,
        ".yml":   _load_yaml,
        ".tex":   _load_latex,
        ".rst":   _load_rst,
        ".md":    _load_text,   # already Markdown
        ".txt":   _load_text,
        ".log":   _load_text,
        ".py":    _load_code,
        ".js":    _load_code,
        ".c":     _load_code,
        ".cpp":   _load_code,
        ".sh":    _load_code,
    }
    fn = dispatch.get(ext)
    if fn is None:
        raise ValueError(f"Unsupported format: '{ext}'")
    return fn(filepath)


# ── Loaders ───────────────────────────────────────────────────────────────────

def _load_pdf(path: str) -> list[Document]:
    from langchain_community.document_loaders import PyPDFLoader
    return PyPDFLoader(path).load()


def _load_text(path: str) -> list[Document]:
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    return [Document(page_content=text, metadata={"source": path, "page": 0})]


def _load_code(path: str) -> list[Document]:
    ext = Path(path).suffix.lstrip(".")
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    md = f"```{ext}\n{text}\n```"
    return [Document(page_content=md, metadata={"source": path, "page": 0})]


def _load_csv(path: str) -> list[Document]:
    import pandas as pd
    sep = "\t" if path.lower().endswith(".tsv") else ","
    df = pd.read_csv(path, sep=sep, dtype=str).fillna("")
    pages, chunk = [], 200
    for i in range(0, max(len(df), 1), chunk):
        block = df.iloc[i : i + chunk]
        try:
            md = block.to_markdown(index=False)
        except Exception:
            md = block.to_string(index=False)
        pages.append(Document(
            page_content=md,
            metadata={"source": path, "page": i // chunk},
        ))
    return pages or [Document(page_content="(empty table)", metadata={"source": path, "page": 0})]


def _load_tsv(path: str) -> list[Document]:
    return _load_csv(path)


def _load_json(path: str) -> list[Document]:
    with open(path, encoding="utf-8", errors="replace") as f:
        data = json.load(f)
    text = json.dumps(data, indent=2, ensure_ascii=False)
    return [Document(page_content=f"```json\n{text}\n```",
                     metadata={"source": path, "page": 0})]


def _load_jsonl(path: str) -> list[Document]:
    docs, page, buffer = [], 0, []
    with open(path, encoding="utf-8", errors="replace") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                buffer.append(json.loads(line))
            except json.JSONDecodeError:
                buffer.append({"_raw": line})
            if len(buffer) >= 50:
                _flush_jsonl(buffer, path, page, docs)
                buffer, page = [], page + 1
    if buffer:
        _flush_jsonl(buffer, path, page, docs)
    return docs or [Document(page_content="(empty)", metadata={"source": path, "page": 0})]


def _flush_jsonl(buffer: list, path: str, page: int, docs: list) -> None:
    text = "\n".join(json.dumps(o, ensure_ascii=False) for o in buffer)
    docs.append(Document(page_content=f"```jsonl\n{text}\n```",
                         metadata={"source": path, "page": page}))


def _load_yaml(path: str) -> list[Document]:
    import yaml
    with open(path, encoding="utf-8", errors="replace") as f:
        data = yaml.safe_load(f)
    text = yaml.dump(data, allow_unicode=True, default_flow_style=False)
    return [Document(page_content=f"```yaml\n{text}\n```",
                     metadata={"source": path, "page": 0})]


def _load_html(path: str) -> list[Document]:
    from bs4 import BeautifulSoup
    html = Path(path).read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "head", "nav", "footer"]):
        tag.decompose()
    md = _html_to_md(str(soup.body or soup))
    return [Document(page_content=_squeeze(md), metadata={"source": path, "page": 0})]


def _load_xml(path: str) -> list[Document]:
    from bs4 import BeautifulSoup
    xml = Path(path).read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(xml, "lxml-xml")
    text = soup.get_text(separator="\n")
    return [Document(page_content=_squeeze(text), metadata={"source": path, "page": 0})]


def _load_docx(path: str) -> list[Document]:
    from docx import Document as DocxDoc
    doc = DocxDoc(path)

    _HEADING = {
        "heading 1": "#", "heading 2": "##", "heading 3": "###",
        "heading 4": "####", "title": "#", "subtitle": "##",
    }

    # Collect body elements in document order via XML iteration
    from docx.oxml.ns import qn
    body = doc.element.body
    lines: list[str] = []

    para_map = {p._element: p for p in doc.paragraphs}
    tbl_map  = {t._element: t for t in doc.tables}

    for child in body:
        if child in para_map:
            para = para_map[child]
            style = para.style.name.lower()
            text  = para.text.strip()
            if not text:
                lines.append("")
                continue
            prefix = _HEADING.get(style, "")
            if prefix:
                lines.append(f"{prefix} {text}")
            else:
                parts = []
                for run in para.runs:
                    t = run.text
                    if run.bold and run.italic:
                        t = f"***{t}***"
                    elif run.bold:
                        t = f"**{t}**"
                    elif run.italic:
                        t = f"*{t}*"
                    parts.append(t)
                lines.append("".join(parts))
        elif child in tbl_map:
            tbl = tbl_map[child]
            if not tbl.rows:
                continue
            header = [c.text.strip() for c in tbl.rows[0].cells]
            lines += [
                "",
                "| " + " | ".join(header) + " |",
                "| " + " | ".join(["---"] * len(header)) + " |",
            ]
            for row in tbl.rows[1:]:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")

    content = _squeeze("\n".join(lines))
    return [Document(page_content=content, metadata={"source": path, "page": 0})]


def _load_doc(path: str) -> list[Document]:
    """Legacy .doc: convert via LibreOffice headless → DOCX → _load_docx."""
    import subprocess
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx",
                 "--outdir", tmpdir, path],
                check=True, capture_output=True, timeout=60,
            )
        except FileNotFoundError:
            raise RuntimeError(
                "LibreOffice is required to load .doc files. "
                "Install it (sudo apt install libreoffice) or convert to .docx."
            )
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"LibreOffice conversion failed: {e.stderr.decode()}")
        out = os.path.join(tmpdir, Path(path).stem + ".docx")
        if not os.path.exists(out):
            raise RuntimeError("LibreOffice produced no output file.")
        return _load_docx(out)


def _load_odt(path: str) -> list[Document]:
    from odf.opendocument import load as odf_load
    from odf import text as odf_text

    doc = odf_load(path)
    lines: list[str] = []

    for elem in doc.text.childNodes:
        tag = getattr(elem, "qname", ("", ""))[1]
        raw = elem.plaintext() if hasattr(elem, "plaintext") else ""
        text = raw.strip()
        if not text:
            lines.append("")
            continue
        if tag == "h":
            level = int(elem.getAttribute("outlinelevel") or 1)
            lines.append("#" * min(level, 6) + f" {text}")
        else:
            lines.append(text)
        lines.append("")

    return [Document(page_content=_squeeze("\n".join(lines)),
                     metadata={"source": path, "page": 0})]


def _load_rtf(path: str) -> list[Document]:
    from striprtf.striprtf import rtf_to_text
    rtf = Path(path).read_text(encoding="utf-8", errors="replace")
    text = rtf_to_text(rtf)
    return [Document(page_content=text, metadata={"source": path, "page": 0})]


def _load_latex(path: str) -> list[Document]:
    from pylatexenc.latex2text import LatexNodes2Text
    tex = Path(path).read_text(encoding="utf-8", errors="replace")
    text = LatexNodes2Text().latex_to_text(tex)
    return [Document(page_content=_squeeze(text), metadata={"source": path, "page": 0})]


def _load_rst(path: str) -> list[Document]:
    rst = Path(path).read_text(encoding="utf-8", errors="replace")
    try:
        from docutils.core import publish_string
        html = publish_string(rst, writer_name="html").decode("utf-8", errors="replace")
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "lxml")
        md = _html_to_md(str(soup.body or soup))
    except ImportError:
        # Minimal RST strip: remove underline-only lines and directives
        md = re.sub(r"^\.\.\s+\S.*$", "", rst, flags=re.MULTILINE)
        md = re.sub(r"^[=\-~^\"#]{3,}$", "", md, flags=re.MULTILINE)
    return [Document(page_content=_squeeze(md), metadata={"source": path, "page": 0})]


def _load_epub(path: str) -> list[Document]:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    book = epub.read_epub(path, options={"ignore_ncx": True})
    docs, page = [], 0
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        html = item.get_content().decode("utf-8", errors="replace")
        soup = BeautifulSoup(html, "lxml")
        for tag in soup(["script", "style"]):
            tag.decompose()
        md = _html_to_md(str(soup.body or soup)).strip()
        if md:
            docs.append(Document(page_content=_squeeze(md),
                                 metadata={"source": path, "page": page}))
            page += 1
    return docs or [Document(page_content="(empty epub)", metadata={"source": path, "page": 0})]


# ── Helpers ───────────────────────────────────────────────────────────────────

def _html_to_md(html: str) -> str:
    """Convert HTML string to Markdown. Uses markdownify if available."""
    try:
        import markdownify
        return markdownify.markdownify(html, heading_style="ATX", strip=["a"])
    except ImportError:
        from bs4 import BeautifulSoup
        return BeautifulSoup(html, "lxml").get_text(separator="\n")


def _squeeze(text: str, max_blanks: int = 2) -> str:
    """Collapse runs of more than max_blanks consecutive blank lines."""
    return re.sub(r"\n{%d,}" % (max_blanks + 1), "\n" * max_blanks, text).strip()
